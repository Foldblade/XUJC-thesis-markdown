# encoding: utf-8
import json
import zipfile
import os
import sys
import shutil
import base64
import re
import argparse
from docx.shared import Cm
from docx import Document
from docx.enum.section import WD_SECTION
from lxml import etree
import progressbar
import requests


WHERE_SCRIPT = os.path.split(os.path.realpath(__file__))[0]


def unzip(file_name):
    """
    解压指定文件
    :param file_name: 需解压的文件名
    :return: 解压后的文件夹路径
    """
    zip_file = zipfile.ZipFile(file_name)
    if os.path.isdir(file_name + "_files"):
        shutil.rmtree(file_name + "_files")
    else:
        os.mkdir(file_name + "_files")
    for names in zip_file.namelist():
        zip_file.extract(names, file_name + "_files/")
    zip_file.close()
    return file_name + "_files"


def zipDir(dirpath, outFullName):
    """
    压缩指定文件夹
    :param dirpath: 目标文件夹路径
    :param outFullName: 压缩文件保存路径
    :return: 无
    """
    zip_file = zipfile.ZipFile(outFullName, "w", zipfile.ZIP_DEFLATED)
    for path, dirnames, filenames in os.walk(dirpath):
        # 去除目标与路径，只对目标文件夹下边的文件及文件夹进行压缩
        fpath = path.replace(dirpath, '')

        for filename in filenames:
            zip_file.write(os.path.join(path, filename),
                           os.path.join(fpath, filename))
    zip_file.close()


def download(url, file_name):
    """
    下载文件
    :param url: 下载链接
    :param file_name: 下载后保存的文件名
    :return: 无
    """
    if not os.path.exists(file_name):
        print(f"Downloading {url} to {file_name}...")
        with open(file_name, 'wb') as f:
            response = requests.request("GET", url, stream=True)
            content_length = int(response.headers.get("Content-Length"))
            widgets = ['Progress: ', progressbar.Percentage(), ' ',
                       progressbar.Bar(marker='#', left='[', right=']'),
                       ' ', progressbar.ETA(), ' ', progressbar.FileTransferSpeed()]
            progress_bar = progressbar.ProgressBar(
                widgets=widgets, maxval=content_length).start()
            for chunk in response.iter_content(chunk_size=1):
                if chunk:
                    f.write(chunk)
                    f.flush()
                progress_bar.update(len(chunk) + 1)
            progress_bar.finish()
    else:
        print("Already exists: %s, skip." % file_name)


def add_cover_image(filepath):
    '''
    添加校徽与校名图片、处理封底
    :param file: 目标文件路径
    :return: 无
    '''

    print("Adding cover image...")
    doc = Document(filepath)
    paragraphs = doc.paragraphs
    found_logo = False
    found_school_name = False
    for p in paragraphs:
        if(p.text == r"{{logo_image}}"):
            p.clear()
            r = p.add_run()
            r.add_picture(os.path.join(
                WHERE_SCRIPT, 'assets/logo_image.png'), width=Cm(5.45), height=Cm(5.45))
            found_logo = True
        elif(p.text == r"{{school_name_image}}"):
            p.clear()
            r = p.add_run()
            r.add_picture(os.path.join(
                WHERE_SCRIPT, 'assets/school_name_image.png'), width=Cm(11.43), height=Cm(1.91))
            found_school_name = True
        if(found_logo and found_school_name):
            break

    # 添加封底空白页
    if ARGS.no_blank_back_cover:
        back_cover = doc.add_section(WD_SECTION.ODD_PAGE)
        header = back_cover.header
        header.is_linked_to_previous = False
        footer = back_cover.footer
        footer.is_linked_to_previous = False
    doc.save(filepath)
    print("Adding cover image done.")


def pandoc_process(*, source=os.path.join(WHERE_SCRIPT, 'demo/paper.md'),
                   output=os.path.join(WHERE_SCRIPT, 'build/pandoc_processed.docx')):
    '''
    将 Markdown 文档通过 Pandoc 转换成 Word 文档
    :return: os.system 包裹的 pandoc 命令，可取返回值
    '''
    print("Using Pandoc to convert Markdown to Word...")
    # 检测 PATH 中是否存在 pandoc 命令，不存在则将 bin 临时加入 PATH
    if shutil.which('pandoc') is None:
        print("Pandoc not found in PATH, temporarily add the bin directory to the PATH...")
        os.environ["PATH"] += os.pathsep + \
            os.path.join(WHERE_SCRIPT, 'bin') + os.pathsep
    pandoc_command = ('pandoc '
                      # 注意：每行后面必须加个空格
                      # docx 样式模板
                      + '--reference-doc "%s" ' % os.path.join(WHERE_SCRIPT, 'assets/template.docx')
                      + '-o "%s" ' % output  # 目标输出文件
                      + ('--metadata-file="%s" ' %
                         ARGS.metadata_file if ARGS.metadata_file is not None else '')  # 元数据文件
                      + ('--bibliography "%s" ' %
                         ARGS.bibliography if ARGS.bibliography is not None else '')  # 引用文件
                      # 引用格式，预处理时会自动下载
                      + '--csl "%s" ' % os.path.join(WHERE_SCRIPT, 'assets/chinese-gb7714-2005-numeric.csl')
                      + '--number-sections '  # 章节自动编号

                      # 资源文件路径，默认与输入文件一致
                      + '--resource-path="%s" ' % os.path.dirname(source)
                      + '--highlight-style monochrome '  # 考虑到最后需要打印，选了个灰度的高亮风格
                      # 图片自动编号
                      + '--filter "%s" ' % os.path.join(
                          WHERE_SCRIPT, 'pandoc_fignos.py')
                      # 表格自动编号
                      + '--filter "%s" ' % os.path.join(
                          WHERE_SCRIPT, 'pandoc_tablenos.py')
                      # 公式自动编号
                      + '--filter "%s" ' % os.path.join(
                          WHERE_SCRIPT, 'pandoc_eqnos.py')

                      # 上面三个自动编号过滤器必须前置于我们的自定义过滤器与 --citeproc

                      # 自定义过滤器
                      + '--filter "%s" ' % os.path.join(WHERE_SCRIPT, 'filter.py')
                      + '--citeproc '  # 处理引用
                      + source)
    print("Pandoc command: ")
    print(pandoc_command)
    print("Here goes with the Pandoc debug: ")
    return os.system(pandoc_command)


def document_process(dir_path):
    '''
    文档 后处理 流程。主要是 XML 操作。
    :param dir_path: 待处理的解压后的 docx 目录
    :return: 无
    '''
    tree = etree.parse(os.path.join(dir_path, 'word/document.xml'))
    root = tree.getroot()

    namespaces = root.nsmap

    body = root.find("w:body", namespaces)

    for p in body.findall("w:p", namespaces):
        pPr = p.find("w:pPr", namespaces)
        sectPr = pPr.findall("w:sectPr", namespaces)
        # 移除第一个分节符本身及其之前的全部内容
        if(not sectPr):
            body.remove(p)
        else:
            body.remove(p)
            break

    for p in body.findall("w:p", namespaces):
        for r in p.findall("w:r", namespaces):
            t = r.find("w:t", namespaces)
            rpr = r.find("w:rPr", namespaces)
            # 处理 Pandoc 生成的 H1 编号为"第 x 章"，并将 tab 换成 空格
            if rpr is not None:
                rStyle = rpr.find("w:rStyle", namespaces)
                if rStyle is not None and rStyle.attrib.get("{%s}val" % namespaces["w"]) == "SectionNumber":
                    # H1 标号数字变 第 x 章
                    if re.match(r'^[0-9]+$', t.text):
                        t.text = "第 %s 章" % t.text
                    # 章节标号后 Tab 换空格
                    maybeTab = r.getnext().find("w:tab", namespaces)
                    if maybeTab is not None:
                        r.getnext().remove(maybeTab)
                        space = etree.SubElement(
                            r.getnext(), "{%s}t" % namespaces["w"])
                        space.text = " "
                        space.set(
                            "{http://www.w3.org/XML/1998/namespace}space", "preserve")
                else:
                    break
            else:
                break

    # 禁止表格行进行跨页换行
    # for tbl in body.findall("w:tbl", namespaces):
    #     for tr in tbl.findall("w:tr", namespaces):
    #         trPr = tr.find("w:trPr", namespaces)
    #         if trPr is not None and trPr.find("w:cantSplit", namespaces) is None:
    #             etree.SubElement(
    #                 trPr, "{%s}cantSplit" % namespaces["w"])
    #         elif trPr is None:
    #             new_trPr = etree.SubElement(
    #                 tr, "{%s}trPr" % namespaces["w"])
    #             etree.SubElement(
    #                 new_trPr, "{%s}cantSplit" % namespaces["w"])

    tree.write(os.path.join(dir_path, 'word/document.xml'),
               encoding='utf-8', xml_declaration=True, standalone=True)


def modify_document_setting(dir_path):
    '''
    后处理流程之二
    修改 Word 文档版式设定中的 字符间距控制 - 只压缩标点符号
    设定 <w:defaultTabStop w:val="420"/>， Tab 宽度 2 字符
    :param dir_path: 待处理的解压后的 docx 目录
    :return: 无
    '''
    tree = etree.parse(os.path.join(dir_path, 'word/settings.xml'))
    root = tree.getroot()

    namespaces = root.nsmap

    root.find("w:characterSpacingControl", namespaces).set(
        '{%s}val' % namespaces["w"], 'compressPunctuation')
    root.find("w:defaultTabStop", namespaces).set(
        '{%s}val' % namespaces["w"], '420')

    tree.write(os.path.join(dir_path, 'word/settings.xml'),
               encoding='utf-8', xml_declaration=True, standalone=True)


def pre_process():
    '''
    预处理流程
    预处理将会下载 cls 文件并生成模板 docx
    :return: 无
    '''
    print("Pre-processing...")
    # 创建 build
    if not os.path.exists(os.path.join(WHERE_SCRIPT, 'build')):
        os.mkdir(os.path.join(WHERE_SCRIPT, 'build'))
    # 创建 bin
    if not os.path.exists(os.path.join(WHERE_SCRIPT, 'bin')):
        os.mkdir(os.path.join(WHERE_SCRIPT, 'bin'))
    download("https://www.zotero.org/styles/chinese-gb7714-2005-numeric",
             os.path.join(WHERE_SCRIPT, 'assets/chinese-gb7714-2005-numeric.csl'))
    download("https://cdn.jsdelivr.net/gh/foldblade/pandoc-fignos@section-separator/pandoc_fignos.py",
             os.path.join(WHERE_SCRIPT, 'pandoc_fignos.py'))
    download("https://cdn.jsdelivr.net/gh/foldblade/pandoc-tablenos@section-separator/pandoc_tablenos.py",
             os.path.join(WHERE_SCRIPT, 'pandoc_tablenos.py'))
    download("https://cdn.jsdelivr.net/gh/foldblade/pandoc-eqnos@section-separator/pandoc_eqnos.py",
             os.path.join(WHERE_SCRIPT, 'pandoc_eqnos.py'))
    if not os.path.exists(os.path.join(WHERE_SCRIPT, 'assets/template.docx')):
        zipDir(os.path.join(WHERE_SCRIPT, 'assets/template'),
               os.path.join(WHERE_SCRIPT, 'assets/template.docx'))
    else:
        print("Already exists: %s, skip." %
              os.path.join(WHERE_SCRIPT, 'assets/template.docx'))

    with open(os.path.join(WHERE_SCRIPT, 'assets/logo_image.b64'), 'r') as f:
        imgdata = base64.b64decode(f.read())
        file = open(os.path.join(WHERE_SCRIPT, 'assets/logo_image.png'), 'wb')
        file.write(imgdata)
        file.close()

    with open(os.path.join(WHERE_SCRIPT, 'assets/school_name_image.b64'), 'r') as f:
        imgdata = base64.b64decode(f.read())
        file = open(os.path.join(
            WHERE_SCRIPT, 'assets/school_name_image.png'), 'wb')
        file.write(imgdata)
        file.close()

    print("Pre-processing done.\n")


def post_process(*, source=os.path.join(WHERE_SCRIPT, 'build/pandoc_processed.docx'),
                 output=os.path.join(WHERE_SCRIPT, 'build/final.docx')):
    '''
    后处理将向过滤器输出的 docx 文件添加校徽与校名图片，
    设置字符间距控制为“只压缩标点符号”
    设定 <w:defaultTabStop w:val="420"/>， Tab 宽度 2 字符
    修改 Header 1 为“第 x 章”
    替换 Header 后面的 Tab 为空格

    部分后处理需要解压缩 docx 文件，大部分是 XML 操作，请参阅 [document_process()]
    :return: 无
    '''
    print("Post-processing...")
    add_cover_image(source)
    unzipped_dir_path = unzip(source)
    print(r"Removing first \newSectionInNewPage and previous contents...")
    document_process(unzipped_dir_path)
    print(r"Modifying document setting...")
    modify_document_setting(unzipped_dir_path)
    if os.path.join(WHERE_SCRIPT, 'build/final.docx') == output:
        zipDir(unzipped_dir_path, output)
    else:
        zipDir(unzipped_dir_path, os.path.join(
            WHERE_SCRIPT, 'build/final.docx'))
        shutil.copy(os.path.join(WHERE_SCRIPT, 'build/final.docx'), output)
    zipDir(unzipped_dir_path, output)
    print("Post-processing done.")
    print(f"Output file: {output}")


def init(path):
    '''
    新建（复制最小化模板到指定位置）
    :return: 无
    '''
    dest = os.path.join(os.getcwd(), path)
    print(f"Initialize scaffold in {dest}...")
    if not os.path.exists(dest):
        shutil.copytree(os.path.join(WHERE_SCRIPT, 'assets/scaffold'), dest)
    else:
        print("The specified directory already exists.")


def clean():
    '''
    清理工作
    :return: 无
    '''
    print("Cleaning...")
    dirs = ['build', 'bin', 'docs_site', 'venv']
    for directory in dirs:
        if os.path.exists(os.path.join(WHERE_SCRIPT, directory)):
            shutil.rmtree(os.path.join(WHERE_SCRIPT, directory))
    files = ['pandoc_eqnos.py', 'pandoc_fignos.py', 'pandoc_tablenos.py',
             './assets/chinese-gb7714-2005-numeric.csl', './assets/logo_image.png',
             './assets/school_name_image.png', './assets/template.docx',
             '.gui_config.json']
    for file in files:
        if os.path.exists(os.path.join(WHERE_SCRIPT, file)):
            os.remove(os.path.join(WHERE_SCRIPT, file))
    print("Cleaning done.")


def split_version(version_string):
    """
    将字符串按照 "." 分割，并将每部分转成数字
    :param version_string: 版本号字符串
    :return: 版本号列表
    """
    version_list = version_string.split('.')
    return [int(n) for n in version_list]


def justify_two_version_list(list1, list2):
    """
    如果两个数字列表长度不一，需要将短一点的列表末尾补零，让它们长度相等
    :param lst1:
    :param lst2:
    :return:
    """
    len1, len2 = len(list1), len(list2)
    if len1 > len2:
        list1 += [0] * (len1-len2)
    elif len1 < len2:
        list2 += [0] * (len2-len1)
    return list1, list2


def compare_version_lists(current_version, to_compair_version):
    """
    比较版本号列表，从高位到底位逐位比较，根据情况判断大小。
    :param current_version: 当前版本
    :param to_compair_version: 要比较的版本
    :return:
    """
    for v1, v2 in zip(current_version, to_compair_version):
        if v1 > v2:
            return False
        elif v1 < v2:
            return True
    return False


def check_update():
    '''
    检查更新
    :return: 存在更新：新版本号，不存在更新：None
    '''
    url = 'https://api.github.com/repos/Foldblade/XUJC-thesis-markdown/releases/latest'

    try:
        response = requests.get(url)
    except Exception as e:
        return None
    else:
        data = response.json()
        tag_name = data['tag_name']
        version1, version2 = justify_two_version_list(
            split_version(VERSION), split_version(tag_name))
        if compare_version_lists(version1, version2):
            return tag_name
        else:
            return None


with open(os.path.join(WHERE_SCRIPT, "VERSION")) as f:
    VERSION = f.read()

parser = argparse.ArgumentParser(
    description="Generate XUJC thesis docx from markdown.")
parser.add_argument("--new", help="Create a new template directory at the specified location. " +
                    "在指定位置新建模板目录。")
pre_post_group = parser.add_mutually_exclusive_group()
pre_post_group.add_argument(
    "--pre", action="store_true", help="Pre processing only. " +
    "仅进行预处理。")
pre_post_group.add_argument(
    "--post", action="store_true", help="Post processing only. " +
    "仅进行后处理。")
parser.add_argument("-F",
                    "--file", help="The file which you want for pandoc convertation or post processing. " +
                    "Pandoc 转换或后处理的文件。")
parser.add_argument("-O",
                    "--output", help="The output file path which you want to save from pandoc convertation or post processing. " +
                    "Pandoc 转换或后处理后文件的保存路径。")
parser.add_argument("-M",
                    "--metadata-file", help="The metadata yaml file path which you want for pandoc convertation." +
                    "用于 Pandoc 转换的 metadata 文件路径。")
parser.add_argument("-B",
                    "--bibliography", help="The Bibtex format bibliography file path which you want for pandoc convertation. " +
                    "用于 Pandoc 转换的 Bibtex 格式参考文献文件路径。")
parser.add_argument("--pandoc-command", help="Overwrite Pandoc command. " +
                    "覆盖 Pandoc 命令。")
parser.add_argument("--no-blank-back-cover", action="store_false", default=True,
                    help="Do not add a blank page as back cover. " +
                    "不要添加空白页作为封底。")
parser.add_argument("--clean", action="store_true", help="Clean the temporary files. " +
                    "清理临时文件。")
parser.add_argument("-V",
                    "--version", action="store_true", help="Check version. " +
                    "查看当前版本。")
# parser.add_argument(
#     "--debug", action="store_true", help="Print debug info. " +
#     "输出调试信息。")

if __name__ == '__main__':
    ARGS = parser.parse_args()
    maybe_update = check_update()

    if ARGS.pre:
        pre_process()
    elif ARGS.post:
        post_process()
    elif ARGS.version:
        print(f"XUJC-thesis-markdown v{VERSION}")
    elif ARGS.new:
        init(ARGS.new)
    elif ARGS.clean:
        clean()
    elif ARGS.pandoc_command is not None:
        pre_process()
        print("Your Pandoc command: ")
        print(ARGS.pandoc_command)
        print("Here goes with the Pandoc debug: ")
        status = os.system(ARGS.pandoc_command)
        if (status == 0):
            print("Pandoc convertation done.\n")
            post_process(output=os.path.join(
                os.getcwd(), ARGS.output))
        else:
            print(
                "Pandoc convertation failed. Please check the Pandoc command and debug info.\n")
    elif not ARGS.pre and \
            not ARGS.post and \
            ARGS.output is not None and \
            ARGS.file is not None and \
            ARGS.metadata_file is not None:
        pre_process()
        status = pandoc_process(source=ARGS.file)
        if (status == 0):
            print("Pandoc convertation done.\n")
            post_process(output=os.path.join(
                os.getcwd(), ARGS.output))
        else:
            print(
                "Pandoc convertation failed. Please check the Pandoc command and debug info.\n")

    if(maybe_update) is not None:
        print(f"The newest version is: v{maybe_update}")
        print(
            "You can download it at: https://github.com/Foldblade/XUJC-thesis-markdown/releases/latest")
